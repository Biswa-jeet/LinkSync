import logging
from youtubesearchpython import VideosSearch
import re
from urllib.parse import urlparse, parse_qs
from docx import Document
from docx2pdf import convert

# Configure logging
logging.basicConfig(filename='youtube_search.log', level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')

def search_youtube(query):
    try:
        videos_search = VideosSearch(query, limit=1)
        results = videos_search.result()['result']
        if results:
            return results[0]['link']
    except Exception as e:
        logging.error(f"An error occurred while searching YouTube: {e}")
    return None

def extract_video_id(youtube_link):
    try:
        parsed_url = urlparse(youtube_link)
        if parsed_url.hostname in ('youtu.be', 'www.youtube.com', 'youtube.com'):
            if parsed_url.path == '/watch':
                query_params = parse_qs(parsed_url.query)
                video_id = query_params.get('v', [])[0]
            else:
                video_id = parsed_url.path.lstrip('/')
            return video_id
    except Exception as e:
        logging.error(f"An error occurred while extracting video ID: {e}")
    return None

def generate_audio_download_link(video_id):
    if video_id:
        return f"https://mp3-convert.org/en1/?v={video_id}"

def generate_video_download_link(video_id):
    if video_id:
        return f"https://www.keepvid.to/317?url=https://www.youtube.com/watch?v={video_id}"

def process_line(line, doc):
    line = line.strip()
    youtube_link = search_youtube(line)
    if youtube_link:
        video_id = extract_video_id(youtube_link)
        if video_id:
            audio_download_link = generate_audio_download_link(video_id)
            video_download_link = generate_video_download_link(video_id)
            if audio_download_link and video_download_link:
                # Add indexed text
                p = doc.add_paragraph()
                p.add_run(line).bold = True  # Add indexed text
                # Add sub-indexed links
                p = doc.add_paragraph(style='BodyText')
                p.add_run(f"Audio: ").italic = True
                p.add_run(audio_download_link)
                p = doc.add_paragraph(style='BodyText')
                p.add_run(f"Video: ").italic = True
                p.add_run(video_download_link)
                logging.info(f"Found YouTube link for '{line}': {youtube_link}")
                logging.info(f"Download audio link: {audio_download_link}")
                logging.info(f"Download video link: {video_download_link}")
            else:
                logging.error(f"Failed to generate download links for '{line}'")
        else:
            logging.warning(f"No video ID extracted from YouTube link for '{line}'")
    else:
        logging.info(f"No YouTube link found for '{line}'")

def validate_file_path(file_path):
    if not file_path:
        raise ValueError("File path cannot be empty")
    return True

def read_txt_file(file_path):
    lines = []
    with open(file_path, 'r') as file:
        lines = file.readlines()
    return lines

def read_docx_file(file_path):
    doc = Document(file_path)
    lines = [p.text for p in doc.paragraphs]
    return lines

def write_docx_file(lines, doc, file_path):
    for line in lines:
        process_line(line, doc)
    doc.save(file_path)

def convert_to_pdf(docx_file):
    pdf_file = docx_file.replace(".docx", ".pdf")
    convert(docx_file, pdf_file)
    logging.info(f"Converted {docx_file} to {pdf_file}")
    print(f"Converted {docx_file} to {pdf_file}")

def main():
    try:
        input_file_path = input("Enter the path of the input file: ").strip()
        output_file_path = input("Enter the path of the output .docx file: ").strip()

        # Validate input file path
        if validate_file_path(input_file_path) and validate_file_path(output_file_path):
            if input_file_path.lower().endswith('.txt'):
                lines = read_txt_file(input_file_path)
            elif input_file_path.lower().endswith('.docx'):
                lines = read_docx_file(input_file_path)
            else:
                raise ValueError("Unsupported input file format")

            # Create a new Document
            doc = Document()
            write_docx_file(lines, doc, output_file_path)
            
            convert_choice = input("Do you want to convert the .docx file to PDF? (yes/no): ").strip().lower()
            if convert_choice == 'yes':
                convert_to_pdf(output_file_path)
            elif convert_choice != 'no':
                print("Invalid choice. Program terminated.")
                logging.error("Invalid choice for conversion. Program terminated.")
    except FileNotFoundError:
        logging.error("Input file not found")
    except PermissionError:
        logging.error("Permission denied to access files")
    except ValueError as ve:
        logging.error(f"Invalid input: {ve}")
    except Exception as e:
        logging.error(f"An unexpected error occurred: {e}")

if __name__ == "__main__":
    main()
